"""python-docx によるWordテキスト抽出。"""

import io

import docx

from .types import OcrResult, PageResult


def extract_docx(file_bytes: bytes, document_id: str) -> OcrResult:
    """Wordからテキストを抽出。文書全体=1ページ。"""
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = []
    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # テーブル
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    text = "\n".join(parts)
    pages = [PageResult(page_number=1, text=text)] if text.strip() else []
    return OcrResult(document_id=document_id, pages=pages)
