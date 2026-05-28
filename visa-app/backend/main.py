"""Codex orchestrator – FastAPI service on Cloud Run."""

import io
import copy
import json
import logging
import mimetypes
import os
import subprocess
import tempfile
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv

load_dotenv()
logging.basicConfig(level=os.environ.get("LOG_LEVEL", "INFO").upper())

from fastapi import FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from google.cloud import firestore, storage
from google.cloud.firestore_v1.base_query import FieldFilter
from google.cloud import run_v2
from pydantic import BaseModel

from extractors.document_models import LoadedDocument
from extractors.document_preprocessor import prepare_documents
from extractors.gemini_pipeline import extract_documents
from extractors.types import ExtractionResult
from application_data import (
    build_display_case_data,
    build_application_data as build_application_data_response,
    load_default_form_definitions,
    load_default_mapping,
)

logger = logging.getLogger(__name__)


def _log_extract_metric(event: str, **fields) -> None:
    safe_fields = {
        key: value
        for key, value in fields.items()
        if value is not None
    }
    logger.info(
        "extract_metric event=%s %s",
        event,
        json.dumps(safe_fields, ensure_ascii=False, sort_keys=True),
    )

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
GCS_BUCKET = os.environ.get("GCS_BUCKET", "visa-codex-mvp-data")
GCP_PROJECT = os.environ.get("GCP_PROJECT", "visa-codex-mvp")
GCP_REGION = os.environ.get("GCP_REGION", "asia-northeast1")
CLOUD_RUN_JOB_NAME = os.environ.get("CLOUD_RUN_JOB_NAME", "codex-runner-job")

# ---------------------------------------------------------------------------
# Clients (initialized lazily on first request via module-level singletons)
# ---------------------------------------------------------------------------
db = firestore.Client(project=GCP_PROJECT)
gcs = storage.Client(project=GCP_PROJECT)
jobs_client = run_v2.JobsClient()

# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------
app = FastAPI(title="codex-orchestrator")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# /api prefix middleware – strips "/api" so the frontend can use "/api/cases"
# in production (where Vite dev-proxy is absent) while the route handlers
# keep their canonical paths ("/cases", "/sessions", etc.).
# ---------------------------------------------------------------------------
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest


class StripApiPrefixMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        if request.url.path.startswith("/api/"):
            # Rewrite /api/cases -> /cases (same as Vite dev proxy)
            new_path = request.url.path[4:]  # strip "/api"
            request.scope["path"] = new_path
        elif request.url.path == "/api":
            request.scope["path"] = "/"
        return await call_next(request)


app.add_middleware(StripApiPrefixMiddleware)


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------
class PromptRequest(BaseModel):
    prompt: str

    @property
    def prompt_stripped(self) -> str:
        return self.prompt.strip()


class CreateCaseRequest(BaseModel):
    application_type: str = "certificate_of_eligibility"
    target_status: str = "engineer_humanities_international"


class UpdateCaseRequest(BaseModel):
    case_data: dict | None = None
    settings: dict | None = None
    field_metadata: dict | None = None
    workflow_state: str | None = None


class ExtractRequest(BaseModel):
    backend: str = "gemini"  # "gemini" | "codex"
    pattern: str = "auto"  # "auto" | "text_only" | "pdf_direct" | "text_and_image"
    scoped: bool = True  # True: スコープ別並列抽出（新方式）、False: 1回呼び出し（後方互換）


def _case_summary(data: dict) -> dict:
    case_data = data.get("case_data") or {}
    applicant = case_data.get("applicant") if isinstance(case_data, dict) else {}
    employer = case_data.get("employer") if isinstance(case_data, dict) else {}
    case_meta = case_data.get("case") if isinstance(case_data, dict) else {}
    applicant_name = (
        data.get("applicant_name_preview")
        or (applicant or {}).get("name_roman")
        or (applicant or {}).get("name_kanji")
        or ""
    )
    employer_name = (employer or {}).get("name") or ""
    display_parts = [part for part in (applicant_name, employer_name) if part]
    display_name = " / ".join(display_parts) or data.get("case_id") or case_meta.get("case_id", "")
    return {
        "case_id": data.get("case_id") or case_meta.get("case_id"),
        "display_name": display_name,
        "applicant_name": applicant_name,
        "employer_name": employer_name,
        "target_status": case_meta.get("target_status", ""),
        "application_type": case_meta.get("application_type", ""),
        "workflow_state": data.get("workflow_state") or case_meta.get("workflow_state", "draft"),
        "created_at": data.get("created_at") or data.get("updated_at") or "",
        "updated_at": data.get("updated_at") or "",
        "applicant_name_preview": applicant_name,
    }


def _deep_merge_dict(base: dict, updates: dict) -> dict:
    for key, value in updates.items():
        if isinstance(base.get(key), dict) and isinstance(value, dict):
            _deep_merge_dict(base[key], value)
        else:
            base[key] = value
    return base


def _merge_extracted_case_data(existing_case_data: dict | None, extracted_case_data: dict) -> dict:
    merged = copy.deepcopy(existing_case_data or {})
    _deep_merge_dict(merged, copy.deepcopy(extracted_case_data or {}))
    return merged


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _make_session_id() -> str:
    return f"sess_{uuid.uuid4().hex[:12]}"


def _make_run_id() -> str:
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    suffix = uuid.uuid4().hex[:4]
    return f"run_{ts}_{suffix}"


def _job_full_name() -> str:
    return f"projects/{GCP_PROJECT}/locations/{GCP_REGION}/jobs/{CLOUD_RUN_JOB_NAME}"


def _upload_prompt(session_id: str, run_id: str, prompt: str) -> str:
    """Upload prompt text to GCS and return the gs:// URI."""
    blob_path = f"sessions/{session_id}/runs/{run_id}/prompt.txt"
    bucket = gcs.bucket(GCS_BUCKET)
    bucket.blob(blob_path).upload_from_string(prompt, content_type="text/plain")
    return f"gs://{GCS_BUCKET}/{blob_path}"


def _launch_job(session_id: str, run_id: str, prompt_gcs_uri: str, firestore_doc_path: str):
    """Launch the Cloud Run Job with execution overrides."""
    overrides = run_v2.types.RunJobRequest.Overrides(
        container_overrides=[
            run_v2.types.RunJobRequest.Overrides.ContainerOverride(
                env=[
                    run_v2.types.EnvVar(name="SESSION_ID", value=session_id),
                    run_v2.types.EnvVar(name="RUN_ID", value=run_id),
                    run_v2.types.EnvVar(name="PROMPT_GCS_URI", value=prompt_gcs_uri),
                    run_v2.types.EnvVar(name="GCS_BUCKET", value=GCS_BUCKET),
                    run_v2.types.EnvVar(name="FIRESTORE_DOC_PATH", value=firestore_doc_path),
                ],
            )
        ],
    )
    request = run_v2.RunJobRequest(name=_job_full_name(), overrides=overrides)
    # Fire-and-forget: the LRO will complete asynchronously.
    jobs_client.run_job(request=request)


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------
@app.post("/sessions")
def create_session(body: PromptRequest):
    if not body.prompt_stripped:
        raise HTTPException(status_code=400, detail="Prompt must not be empty")

    session_id = _make_session_id()
    run_id = _make_run_id()
    now = _now_iso()

    # Upload prompt to GCS
    prompt_gcs_uri = _upload_prompt(session_id, run_id, body.prompt_stripped)

    # Firestore doc paths
    session_ref = db.collection("sessions").document(session_id)
    run_ref = session_ref.collection("runs").document(run_id)
    firestore_doc_path = f"sessions/{session_id}/runs/{run_id}"

    # Write session document
    session_ref.set(
        {
            "session_id": session_id,
            "status": "queued",
            "created_at": now,
            "updated_at": now,
            "latest_run_id": run_id,
            "prompt_preview": body.prompt_stripped[:200],
        }
    )

    # Write run document
    run_ref.set(
        {
            "run_id": run_id,
            "status": "queued",
            "created_at": now,
            "prompt_gcs_uri": prompt_gcs_uri,
        }
    )

    # Launch Cloud Run Job
    try:
        _launch_job(session_id, run_id, prompt_gcs_uri, firestore_doc_path)
    except Exception as exc:
        logger.error("Cloud Run Job launch failed for session %s: %s", session_id, exc)
        # Update status but still return the session so the user can inspect
        run_ref.update({"status": "launch_failed", "error": str(exc)})
        session_ref.update({"status": "launch_failed", "updated_at": _now_iso()})
        return {
            "session_id": session_id,
            "run_id": run_id,
            "status": "launch_failed",
            "error": str(exc),
        }

    return {"session_id": session_id, "run_id": run_id, "status": "running"}


@app.get("/sessions")
def list_sessions():
    query = (
        db.collection("sessions")
        .order_by("created_at", direction=firestore.Query.DESCENDING)
        .limit(20)
    )
    return [doc.to_dict() for doc in query.stream()]


@app.get("/sessions/{session_id}")
def get_session(session_id: str):
    session_doc = db.collection("sessions").document(session_id).get()
    if not session_doc.exists:
        raise HTTPException(status_code=404, detail="Session not found")

    data = session_doc.to_dict()

    # Attach latest run status
    latest_run_id = data.get("latest_run_id")
    if latest_run_id:
        run_doc = (
            db.collection("sessions")
            .document(session_id)
            .collection("runs")
            .document(latest_run_id)
            .get()
        )
        if run_doc.exists:
            data["latest_run"] = run_doc.to_dict()

    return data


@app.get("/sessions/{session_id}/result")
def get_result(session_id: str):
    session_doc = db.collection("sessions").document(session_id).get()
    if not session_doc.exists:
        raise HTTPException(status_code=404, detail="Session not found")

    data = session_doc.to_dict()
    latest_run_id = data.get("latest_run_id")
    if not latest_run_id:
        raise HTTPException(status_code=404, detail="No run found")

    # Try to download last_message from GCS
    blob_path = f"sessions/{session_id}/runs/{latest_run_id}/last_message.txt"
    bucket = gcs.bucket(GCS_BUCKET)
    blob = bucket.blob(blob_path)

    if not blob.exists():
        raise HTTPException(status_code=404, detail="Result not available yet")

    content = blob.download_as_text()
    return {"session_id": session_id, "run_id": latest_run_id, "result": content}


@app.get("/sessions/{session_id}/files")
def list_files(session_id: str):
    """List files in the workspace tarball for the latest run."""
    session_doc = db.collection("sessions").document(session_id).get()
    if not session_doc.exists:
        raise HTTPException(status_code=404, detail="Session not found")

    data = session_doc.to_dict()
    latest_run_id = data.get("latest_run_id")
    if not latest_run_id:
        raise HTTPException(status_code=404, detail="No run found")

    blob_path = f"sessions/{session_id}/runs/{latest_run_id}/workspace.tar.zst"
    bucket = gcs.bucket(GCS_BUCKET)
    blob = bucket.blob(blob_path)

    if not blob.exists():
        raise HTTPException(status_code=404, detail="Workspace archive not found")

    with tempfile.TemporaryDirectory() as tmp:
        archive_path = os.path.join(tmp, "workspace.tar.zst")
        blob.download_to_filename(archive_path)
        extract_dir = os.path.join(tmp, "out")
        os.makedirs(extract_dir)
        subprocess.run(
            ["tar", "--zstd", "-xf", archive_path, "-C", extract_dir],
            check=True,
        )

        files = []
        for p in Path(extract_dir).rglob("*"):
            if p.is_file():
                rel = str(p.relative_to(extract_dir))
                if not rel.startswith(".git/") and not rel.startswith(".git\\"):
                    files.append(rel)

    files.sort()
    return {"files": files}


@app.get("/sessions/{session_id}/files/{file_path:path}")
def download_file(session_id: str, file_path: str):
    """Download a specific file from the workspace tarball."""
    session_doc = db.collection("sessions").document(session_id).get()
    if not session_doc.exists:
        raise HTTPException(status_code=404, detail="Session not found")

    data = session_doc.to_dict()
    latest_run_id = data.get("latest_run_id")
    if not latest_run_id:
        raise HTTPException(status_code=404, detail="No run found")

    blob_path = f"sessions/{session_id}/runs/{latest_run_id}/workspace.tar.zst"
    bucket = gcs.bucket(GCS_BUCKET)
    blob = bucket.blob(blob_path)

    if not blob.exists():
        raise HTTPException(status_code=404, detail="Workspace archive not found")

    # Extract to a temp dir that persists until response is sent
    tmp_dir = tempfile.mkdtemp()
    archive_path = os.path.join(tmp_dir, "workspace.tar.zst")
    blob.download_to_filename(archive_path)
    extract_dir = os.path.join(tmp_dir, "out")
    os.makedirs(extract_dir)
    subprocess.run(
        ["tar", "--zstd", "-xf", archive_path, "-C", extract_dir],
        check=True,
    )

    target = Path(extract_dir) / file_path
    if not target.is_file():
        raise HTTPException(status_code=404, detail="File not found in workspace")

    media_type, _ = mimetypes.guess_type(str(target))
    filename = Path(file_path).name
    return FileResponse(
        path=str(target),
        media_type=media_type or "application/octet-stream",
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Case Endpoints
# ---------------------------------------------------------------------------
@app.post("/cases")
def create_case(body: CreateCaseRequest):
    case_id = f"case_{uuid.uuid4().hex[:12]}"
    now = _now_iso()

    case_data = {
        "schema_version": "1.0",
        "case": {
            "case_id": case_id,
            "application_type": body.application_type,
            "target_status": body.target_status,
            "workflow_state": "draft",
        },
        "applicant": {},
        "entry_plan": {},
        "employer": {},
        "employment": {},
        "proxy": {},
        "receiving_method": {},
    }

    doc = {
        "case_id": case_id,
        "workflow_state": "draft",
        "created_at": now,
        "updated_at": now,
        "case_data": case_data,
        "field_metadata": {},
        "review": {},
        "document_manifest": {"documents": []},
        "extraction_session_id": None,
        "confirmed_at": None,
    }

    db.collection("cases").document(case_id).set(doc)

    return {"case_id": case_id, "workflow_state": "draft", "created_at": now}


@app.get("/cases")
def list_cases(
    limit: int = Query(default=20, ge=1, le=100),
    workflow_state: Optional[str] = Query(default=None),
):
    query = db.collection("cases")
    if workflow_state:
        query = query.where(filter=FieldFilter("workflow_state", "==", workflow_state))
        cases = [_case_summary(doc.to_dict()) for doc in query.stream()]
        return sorted(
            cases,
            key=lambda case: case.get("created_at", ""),
            reverse=True,
        )[:limit]
    cases = [_case_summary(doc.to_dict()) for doc in query.stream()]
    return sorted(
        cases,
        key=lambda case: case.get("created_at", ""),
        reverse=True,
    )[:limit]


@app.get("/cases/{case_id}")
def get_case(case_id: str):
    doc = db.collection("cases").document(case_id).get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")
    data = doc.to_dict()
    if data.get("case_data"):
        data["canonical_case_data"] = copy.deepcopy(data["case_data"])
        data["case_data"] = build_display_case_data(
            data["case_data"],
            data.get("settings"),
        )
    return data


@app.patch("/cases/{case_id}")
def update_case(case_id: str, body: UpdateCaseRequest):
    ref = db.collection("cases").document(case_id)
    doc = ref.get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")

    updates: dict = {"updated_at": _now_iso()}

    if body.case_data is not None:
        updates["case_data"] = body.case_data
    if body.settings is not None:
        updates["settings"] = body.settings
    if body.field_metadata is not None:
        updates["field_metadata"] = body.field_metadata
    if body.workflow_state is not None:
        updates["workflow_state"] = body.workflow_state

    ref.update(updates)

    return ref.get().to_dict()


@app.post("/cases/{case_id}/documents")
async def upload_case_document(
    case_id: str,
    file: UploadFile = File(...),
    document_role: str = Form(default="applicant_document_bundle"),
):
    ref = db.collection("cases").document(case_id)
    doc = ref.get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")

    document_id = f"doc_{uuid.uuid4().hex[:8]}"
    original_filename = file.filename or "upload"
    gcs_path = f"cases/{case_id}/documents/{document_id}_{original_filename}"

    bucket = gcs.bucket(GCS_BUCKET)
    blob = bucket.blob(gcs_path)

    content = await file.read()
    blob.upload_from_string(
        content,
        content_type=file.content_type or "application/octet-stream",
    )

    doc_entry = {
        "document_id": document_id,
        "file_name": original_filename,
        "gcs_path": gcs_path,
        "document_role": document_role,
        "uploaded_at": _now_iso(),
    }

    ref.update(
        {
            "document_manifest.documents": firestore.ArrayUnion([doc_entry]),
            "updated_at": _now_iso(),
        }
    )

    return {
        "document_id": document_id,
        "file_name": original_filename,
        "gcs_path": gcs_path,
        "document_role": document_role,
    }


@app.get("/cases/{case_id}/documents")
def list_case_documents(case_id: str):
    doc = db.collection("cases").document(case_id).get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")
    data = doc.to_dict()
    return data.get("document_manifest", {"documents": []})


@app.get("/cases/{case_id}/documents/{document_id}/url")
def get_document_url(case_id: str, document_id: str):
    doc = db.collection("cases").document(case_id).get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")

    data = doc.to_dict()
    manifest = data.get("document_manifest", {})
    documents = manifest.get("documents", [])

    target_doc = None
    for d in documents:
        if d.get("document_id") == document_id:
            target_doc = d
            break

    if not target_doc:
        raise HTTPException(status_code=404, detail="Document not found in manifest")

    gcs_path = target_doc["gcs_path"]
    bucket = gcs.bucket(GCS_BUCKET)
    blob = bucket.blob(gcs_path)

    try:
        signed_url = blob.generate_signed_url(
            expiration=timedelta(minutes=15), method="GET"
        )
        return {
            "signed_url": signed_url,
            "document_id": document_id,
            "file_name": target_doc.get("file_name"),
        }
    except Exception:
        return {
            "signed_url": None,
            "gcs_path": gcs_path,
            "document_id": document_id,
            "file_name": target_doc.get("file_name"),
            "note": "Signed URL generation failed. Use gcs_path to access the file directly.",
        }


def _find_document_in_manifest(case_id: str, document_id: str) -> tuple[dict, dict]:
    """Lookup a document entry from a case's manifest. Returns (case_data, doc_entry)."""
    doc = db.collection("cases").document(case_id).get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")

    data = doc.to_dict()
    manifest = data.get("document_manifest", {})
    documents = manifest.get("documents", [])

    for d in documents:
        if d.get("document_id") == document_id:
            return data, d

    raise HTTPException(status_code=404, detail="Document not found in manifest")


def _download_document_bytes(gcs_path: str) -> bytes:
    """Download file bytes from GCS."""
    bucket = gcs.bucket(GCS_BUCKET)
    blob = bucket.blob(gcs_path)
    if not blob.exists():
        raise HTTPException(status_code=404, detail="File not found in GCS")
    return blob.download_as_bytes()


def _content_type_for_filename(filename: str) -> str:
    """Return appropriate Content-Type for a filename."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    content_types = {
        "pdf": "application/pdf",
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    return content_types.get(ext, "application/octet-stream")


@app.get("/cases/{case_id}/documents/{document_id}/content")
def get_document_content(case_id: str, document_id: str):
    """GCSからファイルをダウンロードして直接配信（signed URL不要）"""
    _, doc_entry = _find_document_in_manifest(case_id, document_id)
    gcs_path = doc_entry["gcs_path"]
    file_name = doc_entry.get("file_name", "download")

    file_bytes = _download_document_bytes(gcs_path)
    media_type = _content_type_for_filename(file_name)

    return Response(
        content=file_bytes,
        media_type=media_type,
    )


def _docx_to_html(file_bytes: bytes) -> str:
    import docx
    from html import escape
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = ['<div style="font-family:sans-serif;padding:20px;max-width:800px;margin:auto">']
    for para in doc.paragraphs:
        if para.text.strip():
            style = 'font-weight:bold;font-size:1.2em;margin-top:1em' if para.style.name.startswith('Heading') else ''
            parts.append(f'<p style="{style}">{escape(para.text)}</p>')
    for table in doc.tables:
        parts.append('<table style="border-collapse:collapse;width:100%;margin:1em 0">')
        for row in table.rows:
            parts.append('<tr>')
            for cell in row.cells:
                parts.append(f'<td style="border:1px solid #ddd;padding:6px 8px;font-size:13px">{escape(cell.text)}</td>')
            parts.append('</tr>')
        parts.append('</table>')
    parts.append('</div>')
    return '\n'.join(parts)


def _xlsx_to_html(file_bytes: bytes, sheet_name: str | None = None) -> str:
    import openpyxl
    from html import escape
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    sheets = [wb[sheet_name]] if sheet_name else wb.worksheets
    parts = ['<div style="font-family:sans-serif;padding:20px;max-width:1200px;margin:auto">']
    for ws in sheets:
        # 結合セル情報を収集
        merged_map: dict[tuple[int, int], tuple[int, int]] = {}  # (row,col) -> (rowspan,colspan)
        merged_skip: set[tuple[int, int]] = set()
        for rng in ws.merged_cells.ranges:
            r1, c1, r2, c2 = rng.min_row, rng.min_col, rng.max_row, rng.max_col
            merged_map[(r1, c1)] = (r2 - r1 + 1, c2 - c1 + 1)
            for r in range(r1, r2 + 1):
                for c in range(c1, c2 + 1):
                    if (r, c) != (r1, c1):
                        merged_skip.add((r, c))

        parts.append(f'<h3 style="margin:1em 0 0.5em;color:#333">{escape(ws.title)}</h3>')
        parts.append('<table style="border-collapse:collapse;width:100%;margin-bottom:2em">')
        row_idx = 0
        for row in ws.iter_rows(values_only=False):
            row_num = row[0].row if row else 0
            cells = []
            has_content = False
            for cell in row:
                coord = (cell.row, cell.column)
                if coord in merged_skip:
                    continue
                val = str(cell.value) if cell.value is not None else ''
                # @dropdown 等の data_validation 文字列を除去
                if val.startswith('@'):
                    val = ''
                if val:
                    has_content = True
                val_escaped = escape(val)

                # スタイル: ヘッダー行（1行目）は太字+背景色
                if row_num == 1:
                    style = 'border:1px solid #ccc;padding:6px 10px;font-size:13px;font-weight:bold;background:#4a6fa5;color:#fff;white-space:nowrap'
                else:
                    bg = '#f8f9fa' if row_idx % 2 == 0 else '#fff'
                    style = f'border:1px solid #e0e0e0;padding:5px 10px;font-size:13px;background:{bg}'

                # 結合セル属性
                span_attr = ''
                if coord in merged_map:
                    rs, cs = merged_map[coord]
                    if rs > 1:
                        span_attr += f' rowspan="{rs}"'
                    if cs > 1:
                        span_attr += f' colspan="{cs}"'

                cells.append(f'<td style="{style}"{span_attr}>{val_escaped}</td>')
            if has_content:
                parts.append(f'<tr>{"".join(cells)}</tr>')
                row_idx += 1
        parts.append('</table>')
    parts.append('</div>')
    wb.close()
    return '\n'.join(parts)


@app.get("/cases/{case_id}/documents/{document_id}/sheets")
def get_document_sheets(case_id: str, document_id: str):
    """xlsxのシート名一覧を返す。"""
    _, doc_entry = _find_document_in_manifest(case_id, document_id)
    file_name = doc_entry.get("file_name", "")
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    if ext != "xlsx":
        return {"sheets": []}
    import openpyxl
    file_bytes = _download_document_bytes(doc_entry["gcs_path"])
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
    names = wb.sheetnames
    wb.close()
    return {"sheets": names}


@app.get("/cases/{case_id}/documents/{document_id}/preview")
def get_document_preview(
    case_id: str,
    document_id: str,
    sheet: Optional[str] = Query(None, description="シート名（xlsx用）"),
):
    """docx/xlsxをHTML変換して返す。PDF/画像はそのまま返す。"""
    _, doc_entry = _find_document_in_manifest(case_id, document_id)
    gcs_path = doc_entry["gcs_path"]
    file_name = doc_entry.get("file_name", "download")
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    file_bytes = _download_document_bytes(gcs_path)

    if ext == "docx":
        html = _docx_to_html(file_bytes)
        return HTMLResponse(content=html)
    elif ext == "xlsx":
        html = _xlsx_to_html(file_bytes, sheet_name=sheet)
        return HTMLResponse(content=html)
    else:
        # PDF/画像はそのまま配信
        media_type = _content_type_for_filename(file_name)
        return Response(
            content=file_bytes,
            media_type=media_type,
        )


def _extract_with_gemini(
    case_doc: dict,
    pattern: str,
    scoped: bool = True,
    run_id: str | None = None,
) -> ExtractionResult:
    """Run Gemini extraction synchronously.

    Args:
        case_doc: Full case document from Firestore.
        pattern: Extraction pattern ("auto", "text_only", "pdf_direct", "text_and_image").
        scoped: If True, use scope-parallel extraction (new); if False, legacy single-call.
    """
    run_id = run_id or uuid.uuid4().hex[:12]
    documents = case_doc["document_manifest"]["documents"]
    case_data = case_doc.get("case_data", {})
    case_info = case_data.get("case", case_data)
    case_meta = {
        "case_id": case_info.get("case_id", case_doc.get("case_id", "")),
        "application_type": case_info.get("application_type", ""),
        "target_status": case_info.get("target_status", ""),
    }

    logger.info(
        "Gemini extraction started case_id=%s pattern=%s scoped=%s documents=%d",
        case_doc.get("case_id", ""),
        pattern,
        scoped,
        len(documents),
    )
    extraction_started_at = time.monotonic()
    case_id = case_doc.get("case_id", "")
    _log_extract_metric(
        "extraction_start",
        run_id=run_id,
        case_id=case_id,
        pattern=pattern,
        scoped=scoped,
        document_count=len(documents),
    )

    # Download documents from GCS. The extraction pipeline accepts bytes, so
    # GCS stays at the API boundary and eval can reuse the same pipeline.
    def _download_one(doc):
        started_at = time.monotonic()
        blob = gcs.bucket(GCS_BUCKET).blob(doc["gcs_path"])
        content = blob.download_as_bytes()
        file_name = doc["file_name"]
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        _log_extract_metric(
            "document_downloaded",
            run_id=run_id,
            case_id=case_id,
            document_id=doc["document_id"],
            document_role=doc.get("document_role"),
            ext=ext,
            bytes=len(content),
            elapsed_ms=round((time.monotonic() - started_at) * 1000),
        )
        return LoadedDocument(
            document_id=doc["document_id"],
            file_name=file_name,
            document_role=doc.get("document_role", ""),
            content=content,
        )

    with ThreadPoolExecutor(max_workers=min(len(documents), 8)) as pool:
        loaded_documents: list[LoadedDocument] = list(pool.map(_download_one, documents))
    _log_extract_metric(
        "documents_download_complete",
        run_id=run_id,
        case_id=case_id,
        files=len(loaded_documents),
        total_bytes=sum(len(document.content) for document in loaded_documents),
        elapsed_ms=round((time.monotonic() - extraction_started_at) * 1000),
    )
    logger.info(
        "Gemini extraction downloaded documents case_id=%s files=%d total_bytes=%d",
        case_doc.get("case_id", ""),
        len(loaded_documents),
        sum(len(document.content) for document in loaded_documents),
    )

    def log_document_event(event: str, document: LoadedDocument, fields: dict) -> None:
        _log_extract_metric(
            event,
            run_id=run_id,
            case_id=case_id,
            document_id=document.document_id,
            document_role=document.document_role,
            **fields,
        )

    def log_pipeline_event(event: str, fields: dict) -> None:
        _log_extract_metric(event, **fields)

    prepared = prepare_documents(
        loaded_documents,
        event_logger=log_document_event,
    )
    logger.info(
        "Gemini extraction classified documents case_id=%s pdfs=%d text_docs=%d images=%d",
        case_doc.get("case_id", ""),
        len(prepared.pdf_contents),
        len(prepared.text_contents),
        len(prepared.image_entries),
    )
    result = extract_documents(
        case_meta,
        documents,
        loaded_documents,
        prepared,
        pattern=pattern,
        scoped=scoped,
        run_id=run_id,
        case_id=case_id,
        event_logger=log_pipeline_event,
    )
    _log_extract_metric(
        "extraction_complete",
        run_id=run_id,
        case_id=case_id,
        field_metadata_count=len(result.field_metadata),
        review_keys=list(result.review.keys()) if isinstance(result.review, dict) else [],
        elapsed_ms=round((time.monotonic() - extraction_started_at) * 1000),
    )
    return result


@app.post("/cases/{case_id}/extract-stream")
def start_extraction_stream(case_id: str, body: ExtractRequest = ExtractRequest()):
    """SSE streaming extraction endpoint."""
    ref = db.collection("cases").document(case_id)
    doc = ref.get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")
    data = doc.to_dict()
    data.setdefault("case_id", case_id)
    if not data.get("document_manifest", {}).get("documents"):
        raise HTTPException(status_code=400, detail="No documents uploaded")

    def event_stream():
        run_id = uuid.uuid4().hex[:12]
        stream_started_at = time.monotonic()

        def send(event, **kwargs):
            payload = json.dumps(
                {
                    "event": event,
                    "run_id": run_id,
                    "elapsed_ms": round((time.monotonic() - stream_started_at) * 1000),
                    **kwargs,
                },
                ensure_ascii=False,
            )
            return f"data: {payload}\n\n"

        final_state = "extracting"
        try:
            _log_extract_metric(
                "stream_started",
                run_id=run_id,
                case_id=case_id,
                backend=body.backend,
                pattern=body.pattern,
                scoped=body.scoped,
            )
            ref.update({"workflow_state": "extracting", "updated_at": _now_iso()})
            _log_extract_metric(
                "firestore_state_updated",
                run_id=run_id,
                case_id=case_id,
                workflow_state="extracting",
            )
            yield send("progress", phase="downloading", message="ドキュメントを読み込み中...")

            yield send("progress", phase="extracting", message="Gemini APIで抽出中...")
            result = _extract_with_gemini(
                data,
                body.pattern,
                scoped=body.scoped,
                run_id=run_id,
            )

            if not result.display_case_data:
                raise RuntimeError("抽出結果が空です")

            yield send("progress", phase="saving", message="保存中...")
            merged_case_data = _merge_extracted_case_data(
                data.get("case_data"),
                result.display_case_data,
            )

            ref.update({
                "case_data": merged_case_data,
                "review": result.review,
                "field_metadata": result.field_metadata,
                "extraction": {
                    "backend": "gemini",
                    "run_id": run_id,
                    "pattern": body.pattern,
                    "scoped": body.scoped,
                    "completed_at": _now_iso(),
                },
                "workflow_state": "extracted",
                "updated_at": _now_iso(),
            })

            final_state = "extracted"
            _log_extract_metric(
                "firestore_state_updated",
                run_id=run_id,
                case_id=case_id,
                workflow_state="extracted",
            )
            _log_extract_metric(
                "stream_completed",
                run_id=run_id,
                case_id=case_id,
                elapsed_ms=round((time.monotonic() - stream_started_at) * 1000),
            )
            yield send("complete", workflow_state="extracted")

        except Exception as exc:
            logger.error("Stream extraction failed for case %s: %s", case_id, exc)
            final_state = "failed"
            try:
                ref.update({
                    "extraction": {
                        "backend": "gemini",
                        "run_id": run_id,
                        "pattern": body.pattern,
                        "scoped": body.scoped,
                        "failed_at": _now_iso(),
                        "error_type": type(exc).__name__,
                    },
                    "workflow_state": "failed",
                    "updated_at": _now_iso(),
                })
                _log_extract_metric(
                    "firestore_state_updated",
                    run_id=run_id,
                    case_id=case_id,
                    workflow_state="failed",
                )
            except Exception:
                pass
            error_msg = str(exc)
            if "RESOURCE_EXHAUSTED" in error_msg or "429" in error_msg:
                error_msg = "API利用上限に達しました。しばらく待ってから再度お試しください。"
            _log_extract_metric(
                "stream_failed",
                run_id=run_id,
                case_id=case_id,
                error_type=type(exc).__name__,
                elapsed_ms=round((time.monotonic() - stream_started_at) * 1000),
            )
            yield send("error", error=error_msg)
        finally:
            if final_state == "extracting":
                logger.warning("Stream extraction interrupted for case %s", case_id)
                try:
                    ref.update({
                        "extraction": {
                            "backend": "gemini",
                            "run_id": run_id,
                            "pattern": body.pattern,
                            "scoped": body.scoped,
                            "interrupted_at": _now_iso(),
                        },
                        "workflow_state": "failed",
                        "updated_at": _now_iso(),
                    })
                    _log_extract_metric(
                        "stream_interrupted",
                        run_id=run_id,
                        case_id=case_id,
                        elapsed_ms=round((time.monotonic() - stream_started_at) * 1000),
                    )
                except Exception as update_exc:
                    logger.error(
                        "Failed to mark interrupted extraction as failed for case %s: %s",
                        case_id,
                        update_exc,
                    )

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/cases/{case_id}/extract")
def start_extraction(case_id: str, body: ExtractRequest = ExtractRequest()):
    ref = db.collection("cases").document(case_id)
    doc = ref.get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")

    data = doc.to_dict()
    manifest = data.get("document_manifest", {})
    documents = manifest.get("documents", [])

    if not documents:
        raise HTTPException(
            status_code=400, detail="No documents uploaded for this case"
        )

    if body.backend == "gemini":
        return _start_gemini_extraction(case_id, ref, data, body.pattern, body.scoped)

    return _start_codex_extraction(case_id, ref, data, documents)


def _start_gemini_extraction(
    case_id: str, case_ref, case_doc: dict, pattern: str, scoped: bool = True
) -> dict:
    """Gemini backend: synchronous extraction.

    Uses try/finally to guarantee workflow_state is never left as 'extracting',
    even if the Cloud Run instance is shutting down (SIGTERM) or an unexpected
    error occurs after extraction but before the Firestore update.
    """
    now = _now_iso()
    case_ref.update({"workflow_state": "extracting", "updated_at": now})

    final_state = "failed"
    error_msg: str | None = None
    result: ExtractionResult | None = None

    try:
        result = _extract_with_gemini(case_doc, pattern, scoped=scoped)
        merged_case_data = _merge_extracted_case_data(
            case_doc.get("case_data"),
            result.display_case_data,
        )

        case_ref.update(
            {
                "case_data": merged_case_data,
                "review": result.review,
                "field_metadata": result.field_metadata,  # 互換レイヤーで自動生成済み
                "workflow_state": "extracted",
                "updated_at": _now_iso(),
            }
        )
        final_state = "extracted"

    except Exception as exc:
        logger.error("Gemini extraction failed for case %s: %s", case_id, exc)
        error_msg = str(exc)

    finally:
        if final_state != "extracted":
            # Extraction did not complete successfully — ensure Firestore
            # reflects the failure so the case is never stuck in 'extracting'.
            try:
                case_ref.update(
                    {"workflow_state": "failed", "updated_at": _now_iso()}
                )
            except Exception as update_exc:
                logger.error(
                    "Failed to update workflow_state to failed for case %s: %s",
                    case_id, update_exc,
                )

    if final_state == "extracted":
        return {"status": "completed", "workflow_state": "extracted"}
    return {"status": "failed", "error": error_msg or "unknown error"}


def _start_codex_extraction(
    case_id: str, case_ref, data: dict, documents: list[dict]
) -> dict:
    """Codex backend: async extraction via Cloud Run Job."""
    doc_list_text = "\n".join(
        f"- {d['file_name']} (role: {d['document_role']}, gcs: gs://{GCS_BUCKET}/{d['gcs_path']})"
        for d in documents
    )
    prompt = (
        f"You are a visa application document extractor.\n"
        f"Case ID: {case_id}\n"
        f"Application type: {data.get('case_data', {}).get('case', {}).get('application_type', 'unknown')}\n"
        f"Target status: {data.get('case_data', {}).get('case', {}).get('target_status', 'unknown')}\n\n"
        f"Documents:\n{doc_list_text}\n\n"
        f"Extract all relevant fields from these documents.\n"
        f"Output the following JSON files in the generated/ directory:\n"
        f"- generated/case_data.json\n"
        f"- generated/review.json\n"
        f"- generated/field_metadata.json\n"
    )

    session_id = _make_session_id()
    run_id = _make_run_id()
    now = _now_iso()

    prompt_gcs_uri = _upload_prompt(session_id, run_id, prompt)
    firestore_doc_path = f"sessions/{session_id}/runs/{run_id}"

    # Write session document
    session_ref = db.collection("sessions").document(session_id)
    run_ref = session_ref.collection("runs").document(run_id)

    session_ref.set(
        {
            "session_id": session_id,
            "status": "queued",
            "created_at": now,
            "updated_at": now,
            "latest_run_id": run_id,
            "prompt_preview": prompt[:200],
            "linked_case_id": case_id,
        }
    )

    run_ref.set(
        {
            "run_id": run_id,
            "status": "queued",
            "created_at": now,
            "prompt_gcs_uri": prompt_gcs_uri,
        }
    )

    # Update case with extraction session
    case_ref.update(
        {
            "extraction_session_id": session_id,
            "workflow_state": "extracting",
            "updated_at": now,
        }
    )

    # Launch the job
    try:
        _launch_job(session_id, run_id, prompt_gcs_uri, firestore_doc_path)
    except Exception as exc:
        logger.error("Codex job launch failed for case %s, session %s: %s", case_id, session_id, exc)
        run_ref.update({"status": "launch_failed", "error": str(exc)})
        session_ref.update({"status": "launch_failed", "updated_at": _now_iso()})
        case_ref.update(
            {"workflow_state": "failed", "updated_at": _now_iso()}
        )
        return {
            "session_id": session_id,
            "status": "launch_failed",
            "error": str(exc),
        }

    return {"session_id": session_id, "status": "running"}


@app.get("/cases/{case_id}/extraction-status")
def get_extraction_status(case_id: str):
    ref = db.collection("cases").document(case_id)
    doc = ref.get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")

    data = doc.to_dict()
    session_id = data.get("extraction_session_id")
    if not session_id:
        raise HTTPException(
            status_code=404, detail="No extraction session found for this case"
        )

    # Look up session status
    session_doc = db.collection("sessions").document(session_id).get()
    if not session_doc.exists:
        raise HTTPException(status_code=404, detail="Extraction session not found")

    session_data = session_doc.to_dict()
    status = session_data.get("status", "unknown")

    # If completed and case is still extracting, try to harvest results
    if status == "completed" and data.get("workflow_state") == "extracting":
        latest_run_id = session_data.get("latest_run_id")
        if latest_run_id:
            _harvest_extraction_results(case_id, ref, session_id, latest_run_id)

    return {"status": status, "session_id": session_id}


@app.get("/cases/{case_id}/application-data")
def get_application_data(case_id: str):
    """Return generated RASENS input rows for the Chrome extension."""
    doc = db.collection("cases").document(case_id).get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Case not found")

    data = doc.to_dict()
    if not data.get("case_data"):
        raise HTTPException(status_code=400, detail="case_data not found in case document")

    try:
        return build_application_data_response(
            data,
            load_default_mapping(),
            load_default_form_definitions(),
        )
    except ValueError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


def _harvest_extraction_results(
    case_id: str,
    case_ref,
    session_id: str,
    run_id: str,
):
    """Extract case_data.json, review.json, field_metadata.json from workspace tarball."""
    blob_path = f"sessions/{session_id}/runs/{run_id}/workspace.tar.zst"
    bucket = gcs.bucket(GCS_BUCKET)
    blob = bucket.blob(blob_path)

    if not blob.exists():
        return

    with tempfile.TemporaryDirectory() as tmp:
        archive_path = os.path.join(tmp, "workspace.tar.zst")
        blob.download_to_filename(archive_path)
        extract_dir = os.path.join(tmp, "out")
        os.makedirs(extract_dir)
        subprocess.run(
            ["tar", "--zstd", "-xf", archive_path, "-C", extract_dir],
            check=True,
        )

        updates: dict = {"updated_at": _now_iso(), "workflow_state": "extracted"}

        # Look for generated/ files
        extracted_case_data = None
        for name, field in [
            ("case_data.json", "case_data"),
            ("review.json", "review"),
            ("field_metadata.json", "field_metadata"),
        ]:
            # Search for the file in the extracted workspace
            matches = list(Path(extract_dir).rglob(f"generated/{name}"))
            if matches:
                try:
                    content = matches[0].read_text(encoding="utf-8")
                    parsed = json.loads(content)
                    if field == "case_data":
                        extracted_case_data = parsed
                    else:
                        updates[field] = parsed
                except (json.JSONDecodeError, OSError):
                    pass

        if extracted_case_data is not None:
            current_doc = case_ref.get().to_dict() or {}
            updates["case_data"] = _merge_extracted_case_data(
                current_doc.get("case_data"),
                extracted_case_data,
            )

        case_ref.update(updates)


# ---------------------------------------------------------------------------
# Frontend
# ---------------------------------------------------------------------------
# Serve Vite build assets (JS, CSS, images, etc.)
app.mount("/assets", StaticFiles(directory="static/assets"), name="static-assets")


@app.get("/")
def serve_frontend():
    return FileResponse("static/index.html")


@app.get("/{full_path:path}")
def serve_spa(full_path: str):
    """SPA catch-all: return index.html for any unmatched frontend route.

    Requests starting with "api/" should never reach here (they are rewritten
    by StripApiPrefixMiddleware).  If they do, return 404 instead of HTML to
    avoid confusing API clients with an HTML response.

    Static files (e.g. .wasm, .js placed at root by pdfjs-dist) that exist
    in the static/ directory are served directly with the correct MIME type.
    """
    if full_path.startswith("api/") or full_path == "api":
        raise HTTPException(status_code=404, detail="Not found")

    # Serve root-level static files (e.g. jbig2.wasm, openjpeg.wasm) if they exist
    static_path = Path("static") / full_path
    if static_path.is_file() and not full_path.startswith("."):
        mime_overrides = {
            ".wasm": "application/wasm",
            ".js": "application/javascript",
            ".mjs": "application/javascript",
            ".css": "text/css",
            ".svg": "image/svg+xml",
        }
        ext = Path(full_path).suffix.lower()
        media_type = mime_overrides.get(ext) or mimetypes.guess_type(str(static_path))[0]
        return FileResponse(str(static_path), media_type=media_type)

    return FileResponse("static/index.html")
