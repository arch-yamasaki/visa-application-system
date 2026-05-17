"""extractors.docx_text のユニットテスト。python-docx でdocxを動的生成。"""

import io

import docx

from extractors.docx_text import extract_docx
from extractors.types import OcrResult


def _make_docx(paragraphs: list[str] | None = None, tables: list[list[list[str]]] | None = None) -> bytes:
    """テスト用docxをメモリ上で生成。

    Args:
        paragraphs: 段落テキストのリスト
        tables: テーブルデータ [[[cell, ...], ...], ...]
    """
    doc = docx.Document()
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = max(len(r) for r in table_data) if table_data else 0
            if rows == 0 or cols == 0:
                continue
            table = doc.add_table(rows=rows, cols=cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractDocx:
    def test_basic_paragraph_extraction(self):
        data = _make_docx(paragraphs=["Hello World", "Second paragraph"])
        result = extract_docx(data, document_id="doc_d01")

        assert isinstance(result, OcrResult)
        assert result.document_id == "doc_d01"
        assert len(result.pages) == 1
        assert result.pages[0].page_number == 1
        assert "Hello World" in result.pages[0].text
        assert "Second paragraph" in result.pages[0].text

    def test_table_extraction(self):
        data = _make_docx(tables=[[
            ["Name", "Age"],
            ["Taro", "30"],
        ]])
        result = extract_docx(data, document_id="doc_d02")

        assert len(result.pages) == 1
        text = result.pages[0].text
        assert "Name" in text
        assert "Age" in text
        assert "Taro" in text
        assert "30" in text

    def test_paragraphs_and_tables_combined(self):
        data = _make_docx(
            paragraphs=["Header text"],
            tables=[[["Col1", "Col2"], ["Val1", "Val2"]]],
        )
        result = extract_docx(data, document_id="doc_d03")
        text = result.pages[0].text

        assert "Header text" in text
        assert "Col1" in text
        assert "Val1" in text

    def test_empty_docx_returns_no_pages(self):
        data = _make_docx()
        result = extract_docx(data, document_id="doc_d04")
        assert len(result.pages) == 0

    def test_blank_paragraphs_excluded(self):
        data = _make_docx(paragraphs=["", "  ", "Actual content", ""])
        result = extract_docx(data, document_id="doc_d05")

        assert len(result.pages) == 1
        assert "Actual content" in result.pages[0].text

    def test_words_list_is_empty(self):
        """docx抽出ではwordsは空リスト（バウンディングボックスなし）。"""
        data = _make_docx(paragraphs=["test"])
        result = extract_docx(data, document_id="doc_d06")
        assert result.pages[0].words == []

    def test_japanese_content(self):
        data = _make_docx(paragraphs=["株式会社テスト", "東京都千代田区"])
        result = extract_docx(data, document_id="doc_d07")
        text = result.pages[0].text

        assert "株式会社テスト" in text
        assert "東京都千代田区" in text

    def test_multiple_tables(self):
        data = _make_docx(tables=[
            [["A1", "A2"], ["A3", "A4"]],
            [["B1", "B2"], ["B3", "B4"]],
        ])
        result = extract_docx(data, document_id="doc_d08")
        text = result.pages[0].text

        assert "A1" in text
        assert "B1" in text
